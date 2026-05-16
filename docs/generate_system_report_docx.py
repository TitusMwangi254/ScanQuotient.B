#!/usr/bin/env python3
"""
Generate ScanQuotient System Report (Enhanced Edition) as Word (.docx).

Run: python docs/generate_system_report_docx.py
Output: docs/ScanQuotient_System_Report.docx
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_PATH = Path(__file__).resolve().parent / "ScanQuotient_System_Report.docx"
BRAND = RGBColor(0x70, 0x00, 0xFF)


# ── helpers ─────────────────────────────────────────────────────────────

def add_table(doc, headers: list[str], rows: list[list], col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def add_bullets(doc, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code_block(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


def add_note(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run("Note: ")
    r.bold = True
    p.add_run(text).italic = True


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def setup_header_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("ScanQuotient System Report  |  Page ")
    add_page_number_field(fp)


def add_title_page(doc):
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("ScanQuotient")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = BRAND

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("System Technical Report")
    r.bold = True
    r.font.size = Pt(20)

    ed = doc.add_paragraph()
    ed.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ed.add_run("Enhanced Edition").italic = True

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.add_run("Quantifying Risk. Strengthening Security.")

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Document date: {date.today().strftime('%B %d, %Y')}\n")
    meta.add_run("Project path: ScanQuotient.v2/ScanQuotient.B\n")
    meta.add_run("Scanner API version: 2.4.0 (Flask health endpoint)\n")
    meta.add_run("Report engine: deterministic_report_engine.php v5.0\n\n")
    meta.add_run(
        "This document describes architecture, user-facing modules, backend services, "
        "data persistence, and the vulnerability scanner pipeline for academic, "
        "operational, and stakeholder review."
    )

    doc.add_page_break()


def add_toc_manual(doc):
    doc.add_heading("Table of Contents", level=1)
    add_note(doc, "In Microsoft Word: References → Table of Contents → Automatic, to refresh heading page numbers.")
    parts = [
        "1. Executive Summary",
        "2. System Purpose and Scope",
        "3. Architecture and Deployment",
        "4. Frontend — Public Marketing Site",
        "5. Frontend — Authentication and Onboarding",
        "6. Frontend — User Workspace",
        "7. Frontend — Admin Workspace",
        "8. UI/UX and Accessibility",
        "9. Technology Stack",
        "10. Backend Services and APIs",
        "11. Data Storage and Persistence",
        "12. Web Scanner — Inputs, Pipeline, Outputs",
        "13. Risk Scoring Model",
        "14. Reporting Engine (Deterministic + AI)",
        "15. Subscription Tiers and Feature Gates",
        "16. Security Controls",
        "17. Limitations and Assumptions",
        "18. Recommended Future Extensions",
        "Appendix A — Repository Layout",
        "Appendix B — Scanner API Endpoints",
        "Appendix C — PHP Scanner Backend Scripts",
        "Appendix D — Finding Classification Types",
        "Appendix E — Session Variables",
    ]
    add_bullets(doc, parts)
    doc.add_page_break()


# ── content sections ────────────────────────────────────────────────────

def section_executive_summary(doc):
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "ScanQuotient is a web-based vulnerability assessment platform targeting small and medium "
        "enterprises (SMEs), startups, and cybersecurity learners. Users register, authenticate, "
        "submit a target URL, and receive structured findings with severity ratings, evidence blocks, "
        "remediation guidance, and exportable reports (CSV, HTML, PDF, Word-compatible DOC)."
    )
    doc.add_paragraph(
        "The system separates concerns across three runtime layers: (1) a PHP/Apache application "
        "for UI, sessions, billing, and persistence; (2) a Python Flask microservice on port 5000 "
        "for active security testing; and (3) a MySQL database plus filesystem storage for artefacts."
    )
    add_table(
        doc,
        ["Stakeholder", "Primary value"],
        [
            ["Business owner", "Understand exposure without hiring a full pentest team"],
            ["Developer", "Actionable findings on headers, TLS, injection flaws, misconfigurations"],
            ["Student / learner", "Plain-English evidence and structured reports"],
            ["Administrator", "User lifecycle, security logs, subscriptions, support tickets"],
        ],
    )


def section_purpose(doc):
    doc.add_heading("2. System Purpose and Scope", level=1)
    doc.add_heading("2.1 In scope", level=2)
    add_bullets(
        doc,
        [
            "Unauthenticated marketing, legal pages, contact/feedback",
            "User registration with email verification and optional PayPal upgrade",
            "Session-based login with rate limiting, optional 2FA, certificate agreements",
            "Authenticated vulnerability scans against user-supplied http/https URLs",
            "Scan history, grouping, multi-format exports, email share",
            "Freemium / Pro / Enterprise feature differentiation",
            "Admin: users, scan data, site security, payments, feedback, tickets",
        ],
    )
    doc.add_heading("2.2 Out of scope (current build)", level=2)
    add_bullets(
        doc,
        [
            "Authenticated crawling behind login forms (no stored session replay for targets)",
            "Full OWASP Top 10 exhaustive coverage or compliance certification",
            "Continuous monitoring / scheduled scans (not implemented in-repo)",
            "Multi-region SaaS deployment automation",
        ],
    )
    doc.add_heading("2.3 Design philosophy", level=2)
    doc.add_paragraph(
        "The README positions the product as a lightweight website vulnerability analyzer. "
        "The homepage messaging emphasizes democratizing enterprise-grade testing for African SMEs "
        "and educational users, with CVSS-informed scoring and board-ready PDF output."
    )


def section_architecture(doc):
    doc.add_heading("3. Architecture and Deployment", level=1)
    doc.add_heading("3.1 Logical architecture", level=2)
    add_numbers(
        doc,
        [
            "Browser loads PHP pages from XAMPP (e.g. /ScanQuotient.v2/ScanQuotient.B/...).",
            "Authenticated scan UI POSTs JSON to scan_proxy.php (same origin).",
            "scan_proxy.php validates session + URL, forwards body to http://127.0.0.1:5000/scan.",
            "Python scanner returns JSON; UI polls scan_progress.php for stage updates.",
            "save_scan_report.php persists scan_results, writes Storage/Scan_results/*, may call OpenAI.",
            "finding_ai_report.php enriches individual findings (AI + deterministic fallback).",
        ],
    )
    doc.add_heading("3.2 Deployment requirements", level=2)
    add_table(
        doc,
        ["Component", "Requirement"],
        [
            ["Web server", "Apache + PHP 8.x (XAMPP typical on Windows)"],
            ["Database", "MySQL/MariaDB, database name scanquotient.a1"],
            ["Scanner service", "Python 3 with Flask; must run scanner_api.py on port 5000"],
            ["Composer", "dompdf/dompdf for server-side PDF generation"],
            ["Optional", "OPENAI_API_KEY in env or config/secrets.php for AI reports"],
            ["Optional", "PHPMailer + SMTP for registration, 2FA, password reset emails"],
            ["Timezone", "Africa/Nairobi (EAT, UTC+3) used for scans and DB session"],
        ],
    )
    doc.add_heading("3.3 URL routing convention", level=2)
    doc.add_paragraph(
        "The application uses physical PHP paths under Public/ (open) and Private/ (session-guarded). "
        "Hard-coded base path /ScanQuotient.v2/ScanQuotient.B appears in links and redirects."
    )


def section_public_frontend(doc):
    doc.add_heading("4. Frontend — Public Marketing Site", level=1)
    doc.add_paragraph("Primary file: Public/Homepage/PHP/Frontend/Homepage.php")
    doc.add_heading("4.1 Homepage sections (detailed)", level=2)
    add_table(
        doc,
        ["Section", "ID", "Content and behaviour"],
        [
            ["Header", "—", "Sticky glass nav; theme toggle; Help link; Sign In → Login_page_site.php"],
            ["Hero", "#home", "CTA Start Free Scan → Registration; demo anchor to #features; mock score card"],
            ["Features", "#features", "6 cards: SQLi/XSS, headers, SSL/TLS, ports, risk scoring, PDF reports"],
            ["Services", "#services", "3 modules with onclick modals (web, config, ssl/network)"],
            ["Process", "#how-it-works", "3-step timeline: access → scan URL → intelligence delivery"],
            ["Testimonials", "#testimonials", "Three testimonial cards (marketing content)"],
            ["Pricing", "#pricing", "Starter free / Pro $10 / Enterprise $25 with feature bullets"],
            ["Partners", "#trusted-businesses", "CSS marquee of partner names (duplicated for loop)"],
            ["CTA", "—", "Secondary registration push"],
            ["About", "#about-us", "Mission, vision, founder quote, Unsplash dashboard image"],
            ["Contact", "#contact-us", "Phone, email, Nairobi, WhatsApp, social; POST feedback form"],
            ["Footer", "—", "Product/Company/Support links; FAQ, Help, Terms, Privacy"],
        ],
    )
    doc.add_heading("4.2 Supporting public pages", level=2)
    add_table(
        doc,
        ["Page", "File", "Purpose"],
        [
            ["FAQ", "Homepage/PHP/Frontend/FAQ.php", "Accordion Q&A; theme via localStorage"],
            ["Privacy Policy", "Privacy_policy.php", "Legal text; page-scroll-controls.js"],
            ["Terms of Service", "Terms_of_service.php", "Legal text; scroll controls"],
            ["Feedback API", "Backend/submit_customer_feedback.php", "Inserts customer_feedback rows"],
        ],
    )
    doc.add_paragraph("Public pages include security_headers.php: CSP, X-Frame-Options SAMEORIGIN, nosniff, Referrer-Policy, Permissions-Policy, HSTS when HTTPS.")


def section_auth_frontend(doc):
    doc.add_heading("5. Frontend — Authentication and Onboarding", level=1)
    doc.add_heading("5.1 Login flow", level=2)
    add_table(
        doc,
        ["Step", "Component", "Detail"],
        [
            ["1", "Login_page_site.php", "Form POST to login_handler.php; clears old session; no-cache headers"],
            ["2", "Credential verify", "password_verify(); dummy hash if user missing (timing-safe)"],
            ["3", "Account gates", "deleted_at, email_verified, password reset flags, password_expiry"],
            ["4", "2FA (optional)", "6-digit code, 10 min expiry → Login_OTP_verification.php"],
            ["5", "Certificates", "Pending security_certificates must be accepted"],
            ["6", "Session", "regenerate_id; set role, user_id (UID…), profile_photo"],
            ["7", "Redirect", "user → User_dashboard; admin/super_admin → Admin_dashboard"],
        ],
    )
    doc.add_heading("5.2 Registration flow", level=2)
    add_bullets(
        doc,
        [
            "Registration_page.php collects: first/middle/surname, gender, phone, email, recovery email, security Q&A",
            "submit_user_registration_details.php generates UID + 7 alphanumeric user_id (e.g. UIDWRB3O1P)",
            "Random 12-char password emailed via PHPMailer; 6-digit email verification OTP (5 min expiry)",
            "Email_verification.php → verify.php → Registration_completion_site.php",
            "Payment_page.php / PayPal backends for Pro and Enterprise upgrades",
            "Certificate_agreement.php may gate access post-login",
        ],
    )
    doc.add_heading("5.3 Password recovery", level=2)
    add_bullets(
        doc,
        [
            "Forgot_password.php → find_user.php / send_verify.php",
            "Password_reset_page.php with initiate_reset.php and reset_password.php",
            "Forced reset paths: admin-required, expired token, password_expiry column",
        ],
    )
    doc.add_heading("5.4 Help Center", level=2)
    doc.add_paragraph("Public/Help_center/PHP/Frontend/Help_center.php")
    add_bullets(
        doc,
        [
            "Collapsible left nav: Home, Create Ticket, View Previous, Track Ticket, FAQ link",
            "ticket_page_submission.php stores support_tickets; attachments → Storage/Ticket_attachments/",
            "user_ticket_tracking.php for public progress lookup",
            "Admin counterpart: Private/Ticket_support/Admin_ticket_support.php",
        ],
    )


def section_user_frontend(doc):
    doc.add_heading("6. Frontend — User Workspace", level=1)
    doc.add_heading("6.1 User dashboard", level=2)
    doc.add_paragraph("Private/User_dashboard/PHP/Frontend/User_dashboard.php")
    add_bullets(
        doc,
        [
            "Requires role: user, admin, or super_admin",
            "Intelligence carousel: Network, Browser, Security, System (client-side probes via JS)",
            "Sidebar: Dashboard, New Scan, History, Help Center, Account, theme toggle",
            "Action card: Start New Scan, View Scan History",
            "In-page help modal describing each navigation item",
        ],
    )
    doc.add_heading("6.2 Scanner UI (scan.php)", level=2)
    doc.add_paragraph("Private/Web_scanner/PHP/Frontend/scan.php — largest UI surface (~5700+ lines with embedded JS).")
    add_table(
        doc,
        ["UI area", "Capability"],
        [
            ["Target input", "URL field; Start Scan; Cancel; Clear report cache"],
            ["Progress", "Loader, stage label, progress bar via scan_token polling"],
            ["Results", "Severity filters, finding cards, evidence expansion"],
            ["Sidebar (post-scan)", "Risk score analysis, Scan runtime timeline, Detailed report, AI summary"],
            ["Downloads", "CSV, HTML, PDF, DOC AI summary (enterprise)"],
            ["Share", "Email selected artefacts"],
            ["Enterprise", "AI overview page + chatbot (enterprise_ai_overview.php)"],
            ["Notifications", "Bell badge for user alerts"],
        ],
    )
    doc.add_heading("6.3 Scan history and groups", level=2)
    add_bullets(
        doc,
        [
            "historical_scans.php — list/delete/download past scans; risk badges from summary JSON",
            "scan_groups.php — create/rename groups; assign scans via create_scan_group.php, update_scan_group.php",
            "dismantle_group.php — unlink scans and delete group",
        ],
    )
    doc.add_heading("6.4 Account and profile", level=2)
    add_bullets(
        doc,
        [
            "User_subscription.php — plan display aligned with payments table",
            "My_profile.php — profile photo upload to Storage paths",
        ],
    )


def section_admin_frontend(doc):
    doc.add_heading("7. Frontend — Admin Workspace", level=1)
    doc.add_paragraph("Entry: Admin_dashboard.php (roles admin, super_admin).")
    add_table(
        doc,
        ["Module", "Path", "Capabilities"],
        [
            ["Dashboard", "Admin_dashboard.php", "Same intelligence carousel pattern as user dashboard"],
            ["User Management", "User_management/", "List users, User_detail.php edit roles/photos"],
            ["Data Operations", "admin_data_management.php", "Scan oversight, enterprise AI usage events"],
            ["Site Security", "Site_security.php", "Logs, rate limits, certificate CRUD, acceptances"],
            ["Subscriptions", "Subscription_Manager/Payment_accounts.php", "Payment records"],
            ["Feedback", "Feedback/Feedback.php", "customer_feedback moderation (view/trash)"],
            ["Tickets", "Ticket_support/Admin_ticket_support.php", "Respond to support tickets"],
            ["Email", "External Gmail link", "Operational inbox (not embedded MTA)"],
        ],
    )
    add_note(doc, "super_admin vs admin: both use admin dashboard; Site_security filters may differ by role.")


def section_ux(doc):
    doc.add_heading("8. UI/UX and Accessibility", level=1)
    doc.add_heading("8.1 Design system", level=2)
    add_bullets(
        doc,
        [
            "Fonts: Space Grotesk + Inter (homepage); Font Awesome 6.x icons throughout",
            "CSS variables for theming (--primary, --text-muted, glass borders)",
            "Dark/light theme toggles on homepage, login, help, private sidebars (localStorage/cookies)",
            "Consistent header: brand, tagline, avatar, live clock, help, logout",
            "Glassmorphism cards, purple/cyan gradients, responsive grids",
        ],
    )
    doc.add_heading("8.2 UX patterns", level=2)
    add_table(
        doc,
        ["Principle", "Implementation"],
        [
            ["Progressive disclosure", "Scanner report sidebar hidden until scan completes"],
            ["Error tolerance", "User-safe JSON errors from scan_proxy; inline URL validation"],
            ["Learnability", "sidebar-rich-tip tooltips; expand_security_terms in scanner output"],
            ["Feedback", "Toasts (homepage, help), login error box auto-hide 5s, scan progress stages"],
            ["Mobile", "Homepage hamburger; breakpoints 1024px and 768px in homepage.css"],
        ],
    )
    doc.add_heading("8.3 Accessibility", level=2)
    add_bullets(
        doc,
        [
            "html lang=en; aria-label on theme, help, notifications, scroll controls",
            "role=menu on scanner download popovers; role=dialog on admin confirmation modals",
            "alt text on profile images and about-section illustration (with lazy loading)",
            "Keyboard-operable modals (close buttons, overlay click on homepage service modal)",
        ],
    )


def section_tech(doc):
    doc.add_heading("9. Technology Stack", level=1)
    add_table(
        doc,
        ["Layer", "Technologies and versions"],
        [
            ["Application", "PHP 8.x, Apache, session-based MVC-by-folder"],
            ["Scanner", "Python 3.14+, Flask, flask-cors, flask-limiter, requests, certifi"],
            ["Database", "MySQL scanquotient.a1, PDO prepared statements"],
            ["Frontend", "HTML5, CSS3, vanilla JavaScript (no SPA framework)"],
            ["PDF", "dompdf/dompdf ^3.1 via Composer"],
            ["AI", "OpenAI gpt-4o-mini (finding_ai_report.php); configurable timeout/cache"],
            ["Email", "PHPMailer (registration, 2FA, resets)"],
            ["Payments", "PayPal PHP integration (Public/Payment, Private/Payment)"],
            ["Client PDF fallback", "html2canvas + jsPDF vendored in Web_scanner/Javascript/vendor/"],
            ["Evidence", "evidence_engine.py EvidenceFactory for structured proof blocks"],
        ],
    )


def section_backend(doc):
    doc.add_heading("10. Backend Services and APIs", level=1)
    doc.add_heading("10.1 Auth guard", level=2)
    doc.add_paragraph(
        "sq_auth_guard.php defines sq_require_web_scanner_auth(): requires $_SESSION['authenticated'] "
        "=== true and role in (user, admin, super_admin). Returns 401 JSON for API calls."
    )
    doc.add_heading("10.2 Report pipeline", level=2)
    add_numbers(
        doc,
        [
            "Python emits vulnerabilities[] with evidence strings",
            "deterministic_report_engine.php → DeterministicReport::build() (<1ms, no I/O)",
            "finding_ai_report.php may call OpenAI with concurrency slots + ai_report_cache (24h TTL)",
            "save_scan_report.php builds CSV/HTML/PDF/DOC and INSERT scan_results",
            "ensure_pdf_scan.php / prebuild_missing_pdfs.php backfill missing PDFs",
        ],
    )


def section_data(doc):
    doc.add_heading("11. Data Storage and Persistence", level=1)
    doc.add_heading("11.1 Core tables", level=2)
    add_table(
        doc,
        ["Table", "Purpose", "Key columns / notes"],
        [
            ["users", "Accounts", "user_id (UID…), user_name, email, password_hash, role, profile_photo, email_verified, 2FA flags"],
            ["scan_results", "Saved scans", "scan_json LONGTEXT, report_text, pdf/html/csv/doc paths, group_id"],
            ["scan_groups", "Organization", "user_id, name, created_at"],
            ["payments", "Subscriptions", "email, package, account_status, expires_at"],
            ["customer_feedback", "Contact form", "name, email, subject, message, is_viewed, deleted_at"],
            ["support_tickets", "Help desk", "unique_id, status, messages, attachments"],
            ["security_logs", "Audit", "username, event_type, description, IP, user_agent"],
            ["login_rate_limits", "Brute force", "scope user|ip, fail_count, locked_until"],
            ["security_certificates", "Legal gates", "target_type global|role|user, is_active"],
            ["security_certificate_acceptances", "Proof", "certificate_id, user_id, accepted_ip"],
            ["system_server_logs", "Ops", "event_key, level, source, detail_json"],
            ["ai_report_cache", "Performance", "cache_key, report_json, expires_at"],
            ["ai_concurrency_slots", "Throttling", "slot_token, acquired_at"],
            ["enterprise_ai_usage_events", "Analytics", "Enterprise AI feature tracking"],
        ],
    )
    doc.add_heading("11.2 Filesystem", level=2)
    add_table(
        doc,
        ["Path", "Contents"],
        [
            ["Storage/Scan_results/", "{user_id}_{timestamp}_{host}.csv|html|pdf|doc"],
            ["Storage/Public_images/", "page_icon.png, default-avatar.png"],
            ["Storage/Ticket_attachments/", "Uploaded ticket files"],
            ["config/secrets.php (optional, gitignored)", "OPENAI_API_KEY and other secrets"],
        ],
    )


def section_scanner(doc):
    doc.add_heading("12. Web Scanner — Inputs, Pipeline, Outputs", level=1)
    doc.add_heading("12.1 Inputs", level=2)
    add_table(
        doc,
        ["Parameter", "Default", "Description"],
        [
            ["target", "required", "Full http/https URL"],
            ["scan_token", "optional", "UUID for progress tracking"],
            ["enable_port_scan", "true", "Run PortScanner against hostname"],
            ["port_scan_type", "connect", "connect or syn"],
            ["custom_ports", "null", "Optional port list override"],
        ],
    )
    doc.add_heading("12.2 Scan modules (perform_scan order)", level=2)
    add_table(
        doc,
        ["Module", "Progress %", "What it tests"],
        [
            ["Validation", "5", "URL scheme, host, safety checks"],
            ["Connect", "12", "Initial GET, redirect chain, final_url"],
            ["ssl", "20", "Certificate chain, protocols, TLS vulnerabilities"],
            ["server", "26", "Server/X-Powered-By, technology hints"],
            ["headers", "32", "HSTS, CSP, X-Frame-Options, Referrer-Policy, etc."],
            ["cors", "38", "ACAO, credentials, preflight behaviour"],
            ["redirect", "44", "Open redirect parameters"],
            ["files", "50", "~45 sensitive paths (.env, .git, backups, admin UIs)"],
            ["sqli", "58", "Error-based + blind/time payloads on discovered params"],
            ["xss", "66", "14 reflected XSS payloads"],
            ["info", "72", "Error pages, mixed content, version disclosure"],
            ["config", "78", "Cookie flags, HTTPS usage, misconfigs"],
            ["crawl", "84", "Limited same-host URL discovery; per-page header checks"],
            ["ports", "92", "Common ports; Redis/MongoDB/Elasticsearch warnings if open"],
            ["Finalize", "97", "Deduplication, summary, risk score"],
        ],
    )
    doc.add_heading("12.3 Payload inventory (approximate)", level=2)
    add_table(
        doc,
        ["Category", "Count", "Examples"],
        [
            ["SQLi (error-based)", "16", "' OR 1=1--, UNION SELECT, DROP TABLE probes"],
            ["SQLi (blind/time)", "4", "SLEEP(4), WAITFOR DELAY, pg_sleep, BENCHMARK"],
            ["XSS", "14", "<script>, onerror=, javascript:, SVG onload"],
            ["Open redirect", "5", "evil.com, //evil.com, encoded variants"],
            ["Sensitive paths", "45+", ".env, .git/HEAD, wp-config, actuator, dump.sql"],
        ],
    )
    doc.add_heading("12.4 Severity model", level=2)
    add_bullets(
        doc,
        [
            "critical, high, medium, low — contribute to numeric risk score",
            "info, secure — informational; excluded from score calculation",
            "CVSS score optional per finding where applicable",
        ],
    )
    doc.add_heading("12.5 Output JSON shape", level=2)
    add_code_block(
        doc,
        "{\n"
        "  target, timestamp, scan_duration,\n"
        "  ssl, headers, port_scan, server_info, crawler,\n"
        "  vulnerabilities: [{ name, severity, description, evidence,\n"
        "    remediation, cvss_score, what_we_tested, indicates, how_exploited }],\n"
        "  summary: { total_vulnerabilities, severity_breakdown, risk_score, risk_level, ... },\n"
        "  error\n"
        "}",
    )
    doc.add_heading("12.6 Operational limits", level=2)
    add_bullets(
        doc,
        [
            "Global scan timeout: 180 seconds",
            "HTTP request timeout: 10 seconds per request",
            "Flask rate limit: 10 POST /scan per minute per IP",
            "Default limiter: 200/day, 50/hour on scanner API",
        ],
    )


def section_risk(doc):
    doc.add_heading("13. Risk Scoring Model", level=1)
    doc.add_paragraph(
        "Implemented in scanner_api.py after deduplication. Informational and secure severities "
        "do not add points."
    )
    add_code_block(
        doc,
        "raw_score = (critical × 10) + (high × 5) + (medium × 2) + (low × 1)\n"
        "SCORE_CEILING = 106\n"
        "risk_score = min(round((raw_score / SCORE_CEILING) × 100), 100)",
    )
    add_table(
        doc,
        ["risk_score (0–100)", "risk_level"],
        [
            ["≥ 60", "Critical"],
            ["≥ 30", "High"],
            ["≥ 12", "Medium"],
            ["> 0", "Low"],
            ["0", "Secure"],
        ],
    )
    doc.add_paragraph(
        "The scanner also returns risk_score_detail with contributions per severity, weights, "
        "formula_short, and a plain-English opener sentence per band for UI display."
    )


def section_reporting(doc):
    doc.add_heading("14. Reporting Engine (Deterministic + AI)", level=1)
    doc.add_heading("14.1 Deterministic engine v5.0", level=2)
    doc.add_paragraph(
        "File: Private/Web_scanner/PHP/Backend/deterministic_report_engine.php"
    )
    add_bullets(
        doc,
        [
            "Pure PHP, zero network I/O, sub-millisecond per finding",
            "Every field references actual target, parameter, port, header, or payload from evidence",
            "Compatible output shape for finding_ai_report.php normalize pipeline",
            "det_classify() maps finding text to handler category",
        ],
    )
    doc.add_heading("14.2 AI layer (finding_ai_report.php v4.1)", level=2)
    add_table(
        doc,
        ["Setting", "Value"],
        [
            ["Model", "gpt-4o-mini"],
            ["Temperature", "0.15"],
            ["Max tokens", "820"],
            ["Cache TTL", "86400 seconds (ai_report_cache table)"],
            ["Max concurrent", "6 slots (ai_concurrency_slots)"],
            ["Fallback", "DeterministicReport::build() on timeout or missing API key"],
        ],
    )
    doc.add_heading("14.3 Export formats", level=2)
    add_table(
        doc,
        ["Format", "Generator", "Notes"],
        [
            ["CSV", "build_csv_report()", "Spreadsheet-friendly findings list"],
            ["HTML", "build_html_report()", "Interactive review in browser"],
            ["PDF", "Dompdf from HTML", "Server-side; client jsPDF fallback if needed"],
            ["DOC", "HTML saved as .doc", "Word-compatible for meetings"],
        ],
    )


def section_tiers(doc):
    doc.add_heading("15. Subscription Tiers and Feature Gates", level=1)
    add_table(
        doc,
        ["Tier", "Storage", "AI / advanced"],
        [
            ["Freemium (Starter)", "Max 5 saved scans (scan_results COUNT per user_id)", "Basic scan + reports; upgrade prompt after limit"],
            ["Pro ($10/mo)", "Unlimited saves (per marketing)", "Priority queue (marketing); PDF/CSV export"],
            ["Enterprise ($25/mo)", "Unlimited + compliance messaging", "AI per-finding reports, detailed AI, enterprise overview/chat, DOC AI summary"],
        ],
    )
    doc.add_paragraph(
        "Package resolution: get_user_package.php and save_scan_report.php check payments table "
        "(active, non-expired) by user email, then fall back to users.subscription_plan."
    )


def section_security(doc):
    doc.add_heading("16. Security Controls", level=1)
    add_table(
        doc,
        ["Control", "Detail"],
        [
            ["Login rate limit (user)", "5 failures → 5 minute lock; 10 min rolling window"],
            ["Login rate limit (IP)", "20 failures → 15 minute lock; 15 min rolling window"],
            ["Ambiguous errors", "Same message for bad username vs password"],
            ["Session fixation", "session_regenerate_id(true) on success"],
            ["Password storage", "password_hash / password_verify (bcrypt)"],
            ["2FA", "Optional email OTP, 10 minute expiry"],
            ["Email verification", "Required before full access; 6-digit OTP"],
            ["Certificate gating", "security_certificates targeted by global/role/user"],
            ["Scanner proxy", "No direct browser→Python; PHP validates URL scheme"],
            ["Error hygiene", "display_errors off in scan_proxy; generic client messages"],
            ["Security logging", "security_logs + system_server_logs"],
        ],
    )


def section_limits(doc):
    doc.add_heading("17. Limitations and Assumptions", level=1)
    add_bullets(
        doc,
        [
            "Active scanning only — no passive network monitoring",
            "Heuristic SQLi/XSS — may miss blind vulns or flag WAF-blocked noise",
            "verify=False on scanner HTTP client — correct for assessment tool, risky if mis-deployed",
            "Python API binds 0.0.0.0:5000 with CORS * — must not be internet-exposed without proxy auth",
            "Secrets in PHP source paths (DB root, hardcoded vendor paths) — not production-hardened",
            "Legal: users must only scan systems they own or have written permission to test",
            "AI quality varies with API availability; deterministic layer always available",
        ],
    )


def section_future(doc):
    doc.add_heading("18. Recommended Future Extensions", level=1)
    add_numbers(
        doc,
        [
            "Authenticated scan profiles (cookies, Bearer tokens, form login macros)",
            "Scheduled scans + diff/regression alerts via email",
            "CI/CD plugins (GitHub Action, Jenkins) with fail-on-critical threshold",
            "Organization/team tenancy with shared scan libraries",
            "OWASP ASVS / PCI control mapping in Enterprise PDF templates",
            "Containerized scanner with mTLS to PHP layer",
            "Central secrets via .env; remove hardcoded Dompdf/PHPMailer paths",
            "CSRF, SSRF, JWT, and dependency (SCA) modules",
            "Remediation workflow (assign owner, status, due date)",
            "Metrics dashboard from system_server_logs and scan volume",
            "Swahili/localization for East African market",
            "Public REST API with API keys per subscription tier",
        ],
    )


def appendix_layout(doc):
    doc.add_heading("Appendix A — Repository Layout", level=1)
    add_code_block(
        doc,
        "ScanQuotient.B/\n"
        "  Public/           # Marketing, login, register, help, payment\n"
        "  Private/          # Dashboards, scanner, admin modules\n"
        "    Web_scanner/    # scan.php, PHP backends, CSS, JS\n"
        "    Python_Module/  # scanner_api.py, evidence_engine.py\n"
        "    Admin_dashboard/, User_dashboard/, User_management/\n"
        "    Site_security/, Ticket_support/, Feedback/\n"
        "    Subscription_Manager/, Payment/, My_profile/\n"
        "  Storage/          # Scan_results, images, ticket files\n"
        "  vendor/           # Composer (dompdf)\n"
        "  docs/             # This report generator\n"
        "  config/           # Optional secrets.php (gitignored)",
    )


def appendix_scanner_api(doc):
    doc.add_heading("Appendix B — Scanner API Endpoints", level=1)
    add_table(
        doc,
        ["Endpoint", "Method", "Rate limit", "Purpose"],
        [
            ["/scan", "POST", "10/minute", "Full security scan"],
            ["/scan-progress", "GET", "120/minute", "Poll progress by token"],
            ["/port-scan", "POST", "5/minute", "Standalone port scan"],
            ["/health", "GET", "—", "Status, version 2.4.0, feature list"],
        ],
    )
    doc.add_paragraph("Health endpoint features (v2.4.0) include: ssl_scan, header_analysis, sqli_test, xss_test, port_scan, cors_check, open_redirect, sensitive_files, blind_sqli, mixed_content, secret_detection, deduplication, full_request_evidence, cors_preflight_test, dom_xss_source_context, sql_baseline_comparison, and more.")


def appendix_php_backends(doc):
    doc.add_heading("Appendix C — PHP Scanner Backend Scripts", level=1)
    scripts = [
        "scan_proxy.php — Forward to Python /scan",
        "scan_progress.php — Progress polling",
        "save_scan_report.php — Persist + exports + OpenAI summary",
        "finding_ai_report.php — Per-finding AI/deterministic reports",
        "deterministic_report_engine.php — Fast template engine",
        "enterprise_ai_api.php — Enterprise overview/chat",
        "get_user_package.php — Tier resolution",
        "download_scan.php / share_scan.php — Distribution",
        "delete_scan.php — Remove scan + files",
        "create_scan_group.php / update_scan_group.php / dismantle_group.php",
        "scan_run_timeline.php — Host-scoped history for UI",
        "ensure_pdf_scan.php / prebuild_missing_pdfs.php / missing_pdfs_count.php",
        "upload_pdf.php — Client-generated PDF upload path",
        "track_finding_report_event.php / track_enterprise_ai_event.php — Analytics",
    ]
    add_bullets(doc, scripts)


def appendix_finding_types(doc):
    doc.add_heading("Appendix D — Finding Classification Types", level=1)
    doc.add_paragraph("det_classify() categories handled by DeterministicReport:")
    add_bullets(
        doc,
        [
            "sqli, xss, cors, csrf, redirect, cookie",
            "tls, hsts, csp, clickjacking, mime, header (generic)",
            "port, file, secret, info, generic",
        ],
    )
    doc.add_paragraph(
        "Engine header comment lists coverage: SSL/TLS (6 sub-types), security headers (7), "
        "CORS (3), SQLi (2), XSS (2), open redirect, sensitive files, info disclosure (3), "
        "cookie flags (3), port scan (generic + 4 services)."
    )


def appendix_session(doc):
    doc.add_heading("Appendix E — Session Variables", level=1)
    add_table(
        doc,
        ["Variable", "When set", "Purpose"],
        [
            ["authenticated", "Login success", "Gate private pages"],
            ["role", "Login success", "user | admin | super_admin"],
            ["user_id", "Login success", "Public UID (UIDxxxxxxx)"],
            ["user_pk", "Login success", "Numeric users.id"],
            ["user_name, profile_photo", "Login", "UI display"],
            ["login_time, ip_address, user_agent", "Login", "Session metadata"],
            ["2fa_pending, 2fa_code, 2fa_expires", "2FA flow", "OTP verification"],
            ["auth_mode, pending_agreements", "Gates", "email_verification, certificates"],
            ["loginError", "Failed login", "One-time message on login page"],
            ["feedback_status, toast_message", "Forms", "Flash notifications"],
        ],
    )


def build_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    setup_header_footer(doc)
    add_title_page(doc)
    add_toc_manual(doc)

    section_executive_summary(doc)
    section_purpose(doc)
    section_architecture(doc)
    section_public_frontend(doc)
    section_auth_frontend(doc)
    section_user_frontend(doc)
    section_admin_frontend(doc)
    section_ux(doc)
    section_tech(doc)
    section_backend(doc)
    section_data(doc)
    section_scanner(doc)
    section_risk(doc)
    section_reporting(doc)
    section_tiers(doc)
    section_security(doc)
    section_limits(doc)
    section_future(doc)

    doc.add_page_break()
    appendix_layout(doc)
    appendix_scanner_api(doc)
    appendix_php_backends(doc)
    appendix_finding_types(doc)
    appendix_session(doc)

    doc.add_heading("Document Revision History", level=1)
    add_table(
        doc,
        ["Version", "Date", "Changes"],
        [
            ["1.0", "Initial", "Base system report"],
            ["2.0 Enhanced", str(date.today()), "Expanded sections, appendices, risk formula, security controls, API inventory"],
        ],
    )

    return doc


def main():
    doc = build_document()
    doc.save(OUT_PATH)
    size_kb = OUT_PATH.stat().st_size // 1024
    print(f"Saved: {OUT_PATH} ({size_kb} KB)")


if __name__ == "__main__":
    main()
